# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT_DIR / "case_test_data"
ASSET_DIR = OUT_DIR / "assets"
RENDERED_DIR = OUT_DIR / "rendered"

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]

ACCENT = "2563EB"
ACCENT_DARK = "0B2545"
MUTED = "667085"
GRID = "D0D5DD"
SOFT_BLUE = "EAF2FF"
SOFT_GREEN = "ECFDF3"
SOFT_ORANGE = "FFF4E5"
SOFT_RED = "FEF3F2"
SOFT_GRAY = "F8FAFC"


CASES = [
    {
        "filename": "ds_queue_linear_queue_case.docx",
        "title": "数据结构教学案例：队列的入队与出队",
        "topic": "队列",
        "course": "数据结构",
        "difficulty": "easy",
        "duration": "1 课时 / 45 分钟",
        "keywords": ["队列", "Queue", "FIFO", "入队", "出队", "顺序队列", "链式队列"],
        "summary": "以银行取号和窗口叫号为情境，帮助学生把 FIFO 规则、front/rear 指针变化与顺序队列代码建立对应关系。",
        "scenario": [
            "银行大厅中，客户按取号顺序等待办理业务。新客户到达时进入队尾，柜员办理完成后从队头叫号。这个场景天然对应队列 Queue 的先进先出 FIFO 规则。",
            "课堂中把学生分成“客户”“取号机”“柜员”三个角色，用队伍移动模拟 enqueue 与 dequeue。教师同步在黑板上记录 front、rear 与数组下标的变化，让生活经验和抽象变量一一对应。",
        ],
        "driving_question": "为什么“先到先服务”的场景不能用栈来处理？当队列为空或已满时，程序应该如何给出安全反馈？",
        "objectives": [
            "能准确描述队列的逻辑结构和 FIFO 特性，并能举出至少两个真实应用场景。",
            "能区分队头 front、队尾 rear、入队 enqueue、出队 dequeue 的含义。",
            "能用数组或链表实现基本队列，并解释队空、队满判断条件。",
        ],
        "focus": "FIFO 规则、front/rear 的含义、入队与出队状态变化。",
        "difficulty_point": "把“人排队”的直觉迁移为数组下标与边界判断，避免 front/rear 更新顺序写反。",
        "activities": [
            ("导入", "播放银行叫号或食堂排队图片，提出“如何保证公平处理请求”。", "说出生活中的 FIFO 场景。", "能够主动提到先到先服务。"),
            ("建模", "把队伍映射为数组 data，标注 front 与 rear。", "用座位卡片模拟 A、B、C 入队。", "能指出队头和队尾。"),
            ("操作演示", "执行 enqueue(A)、enqueue(B)、dequeue()、enqueue(C)。", "记录每一步数组内容和指针变化。", "状态表填写完整。"),
            ("代码迁移", "展示伪代码并强调越界保护。", "补全 isEmpty 与 isFull。", "判断条件正确。"),
            ("小结", "对比队列与栈的规则差异。", "用一句话解释 FIFO。", "能用关键词命中本案例。"),
        ],
        "operation_steps": [
            "入队前先判断队列是否已满，避免 rear 越界。",
            "将新元素放入 rear 指向的位置，然后 rear 后移。",
            "出队前先判断队列是否为空，避免读取无效元素。",
            "取出 front 指向的元素，然后 front 后移。",
        ],
        "demo_rows": [
            ["初始", "[]", "0", "0", "队列为空"],
            ["enqueue(A)", "[A]", "0", "1", "A 进入队尾"],
            ["enqueue(B)", "[A, B]", "0", "2", "B 进入队尾"],
            ["dequeue()", "[B]", "1", "2", "A 从队头离开"],
            ["enqueue(C)", "[B, C]", "1", "3", "C 进入队尾"],
        ],
        "misconceptions": [
            "把队列误认为“谁新来谁先处理”，与栈的 LIFO 混淆。",
            "出队后忘记移动 front，导致同一元素被重复读取。",
            "只检查 rear 是否越界，却没有先判断队空或队满。",
        ],
        "exercise": "实现一个顺序队列，支持 enqueue、dequeue、isEmpty 三个操作，并说明当 rear 到达数组末尾时会出现什么问题。",
        "extension": "让学生思考：如果业务窗口有多个，普通队列是否足够？是否需要优先队列或多队列调度？",
        "assessment": "学生能根据操作序列写出队列变化过程，并在代码中正确维护 front 和 rear，即视为达成核心目标。",
        "code": """enqueue(x):
    if isFull(): report overflow
    data[rear] = x
    rear = rear + 1

dequeue():
    if isEmpty(): report underflow
    x = data[front]
    front = front + 1
    return x""",
        "diagram": "queue",
    },
    {
        "filename": "ds_circular_queue_case.docx",
        "title": "数据结构教学案例：循环队列与假溢出问题",
        "topic": "循环队列",
        "course": "数据结构",
        "difficulty": "medium",
        "duration": "1 课时 / 45 分钟",
        "keywords": ["循环队列", "Circular Queue", "队列", "假溢出", "front", "rear", "取模运算"],
        "summary": "以打印任务缓冲区为情境，解释顺序队列的假溢出，并通过取模运算完成数组空间复用。",
        "scenario": [
            "打印机不断接收打印任务，也不断完成并移除任务。如果只让 rear 单向增长，即使数组前部已经空出，也可能因为 rear 到达末尾而拒绝新任务，这就是假溢出。",
            "循环队列把线性数组想象成一个环，rear 与 front 通过取模运算回到数组开头，使已经释放的位置重新可用。",
        ],
        "driving_question": "数组前面明明有空位，为什么顺序队列还会提示“满”？循环队列如何让 rear 回到起点？",
        "objectives": [
            "能解释顺序队列产生假溢出的原因。",
            "能写出循环队列 rear = (rear + 1) % maxSize 的更新逻辑。",
            "能判断循环队列的队空和队满条件，并说明牺牲一个存储单元方案的意义。",
        ],
        "focus": "假溢出、取模运算、front/rear 环形移动。",
        "difficulty_point": "理解 (rear + 1) % maxSize == front 表示队满，而不是简单地 rear == maxSize。",
        "activities": [
            ("问题暴露", "用长度为 5 的数组演示顺序队列假溢出。", "指出数组前部空位为何没有被复用。", "能说出“rear 单向增长”。"),
            ("模型转换", "把数组画成环形，标注 0 到 5 的下标。", "用箭头模拟 rear 回绕。", "能写出取模表达式。"),
            ("条件比较", "对比牺牲一个单元与增加 size 变量两种判满方式。", "选择一种方案并说明理由。", "队空队满条件清晰。"),
            ("代码落地", "给出 enqueue/dequeue 伪代码。", "补全判空与判满。", "边界用例通过。"),
        ],
        "operation_steps": [
            "入队时先判断 (rear + 1) % maxSize == front 是否成立。",
            "若未满，将元素写入 rear 指向的槽位。",
            "rear 更新为 (rear + 1) % maxSize，实现回绕。",
            "出队时读取 front 指向元素，再令 front = (front + 1) % maxSize。",
        ],
        "demo_rows": [
            ["入队 1,2,3,4", "front=0, rear=4", "[1,2,3,4,_]", "rear 指向下一个空位"],
            ["出队两次", "front=2, rear=4", "[_,_,3,4,_]", "数组前部释放"],
            ["入队 5", "front=2, rear=0", "[_,_,3,4,5]", "rear 取模回绕"],
            ["入队 6", "front=2, rear=1", "[6,_,3,4,5]", "复用 0 号槽"],
        ],
        "misconceptions": [
            "认为 rear 到达数组末尾就一定队满。",
            "忽略牺牲一个单元方案，导致队空和队满条件都写成 front == rear。",
            "取模时使用 maxSize - 1，造成下标范围错误。",
        ],
        "exercise": "给定 maxSize=6，依次执行入队 1、2、3、4，出队两次，再入队 5、6、7，请画出循环队列数组状态。",
        "extension": "引导学生讨论操作系统中的环形缓冲区、键盘输入缓冲区和网络包接收队列。",
        "assessment": "学生能正确使用取模表达式维护 front 和 rear，并能说明队满条件，即达到教学目标。",
        "code": """isFull():
    return (rear + 1) % maxSize == front

enqueue(x):
    if isFull(): report overflow
    data[rear] = x
    rear = (rear + 1) % maxSize

dequeue():
    if front == rear: report underflow
    x = data[front]
    front = (front + 1) % maxSize
    return x""",
        "diagram": "circular_queue",
    },
    {
        "filename": "ds_stack_expression_case.docx",
        "title": "数据结构教学案例：栈与表达式括号匹配",
        "topic": "栈",
        "course": "数据结构",
        "difficulty": "easy",
        "duration": "1 课时 / 45 分钟",
        "keywords": ["栈", "Stack", "LIFO", "入栈", "出栈", "括号匹配", "表达式求值"],
        "summary": "通过表达式括号匹配任务，帮助学生理解栈的后进先出 LIFO 特性，以及编译器语法检查中的典型应用。",
        "scenario": [
            "编译器检查代码时，需要判断圆括号、方括号和花括号是否正确匹配。最近出现的左括号必须最先与右括号配对，这正是栈 Stack 的 LIFO 特性。",
            "课堂以表达式 {[a+b]*(c-d)} 为主线，让学生逐字符扫描，遇到左括号压栈，遇到右括号弹栈比较类型。",
        ],
        "driving_question": "为什么括号匹配必须关注“最近出现的左括号”？如果栈空时遇到右括号，说明什么错误？",
        "objectives": [
            "能说明栈的后进先出 LIFO 规则。",
            "能用 push、pop、peek 描述栈的基本操作。",
            "能使用栈完成括号匹配算法设计，并处理不匹配、缺失右括号等异常。",
        ],
        "focus": "LIFO 规则、push/pop 操作、括号类型匹配。",
        "difficulty_point": "把“最近的左括号先匹配”转换为栈顶元素检查，并覆盖栈空、栈非空两类异常。",
        "activities": [
            ("生活类比", "用一摞书演示入栈和出栈顺序。", "预测最先拿走哪本书。", "能说出后进先出。"),
            ("算法模拟", "逐字符扫描表达式 {[a+b]*(c-d)}。", "记录栈内容变化。", "每个右括号都能找到对应左括号。"),
            ("错误诊断", "给出 ([a+b]*{c-d})]。", "定位多出的右括号。", "能说明栈空异常。"),
            ("代码实现", "展示伪代码框架。", "补全 match 函数。", "三类括号都能处理。"),
        ],
        "operation_steps": [
            "扫描表达式中的每一个字符。",
            "遇到左括号时执行 push，将其压入栈顶。",
            "遇到右括号时，若栈空则立即判定不匹配。",
            "若栈非空，pop 栈顶左括号，并检查类型是否对应。",
            "扫描结束后，若栈仍非空，说明存在未闭合的左括号。",
        ],
        "demo_rows": [
            ["{", "push {", "{", "等待右花括号"],
            ["[", "push [", "{ [", "嵌套结构"],
            ["]", "pop [", "{", "方括号匹配"],
            ["(", "push (", "{ (", "进入圆括号"],
            [")", "pop (", "{", "圆括号匹配"],
            ["}", "pop {", "空", "全部匹配"],
        ],
        "misconceptions": [
            "只统计左右括号数量相等，却不检查嵌套顺序。",
            "遇到右括号时没有先判断栈是否为空。",
            "扫描结束后忘记检查栈中是否还有未匹配的左括号。",
        ],
        "exercise": "判断表达式 ([a+b]*{c-d})] 是否括号匹配，并给出栈状态变化过程。",
        "extension": "进一步引出表达式求值、中缀转后缀表达式和函数调用栈。",
        "assessment": "学生能在遇到左括号时入栈、遇到右括号时出栈并检查类型，即视为掌握栈的核心应用。",
        "code": """for each char in expression:
    if char is left bracket:
        push(char)
    else if char is right bracket:
        if stack is empty: return false
        left = pop()
        if not match(left, char): return false

return stack is empty""",
        "diagram": "stack",
    },
    {
        "filename": "ds_singly_linked_list_case.docx",
        "title": "数据结构教学案例：单链表的插入与删除",
        "topic": "单链表",
        "course": "数据结构",
        "difficulty": "medium",
        "duration": "2 课时 / 90 分钟",
        "keywords": ["单链表", "Linked List", "结点", "指针", "插入", "删除", "头结点"],
        "summary": "以通讯录联系人管理为情境，讲解单链表结点结构、next 指针、头结点以及插入删除时的指针修改顺序。",
        "scenario": [
            "通讯录中的联系人数量经常变化。相比数组，单链表在插入和删除联系人时不需要整体移动大量元素，只需要修改相关结点的 next 指针。",
            "课堂中用卡片代表结点，卡片左半写数据域，右半画 next 指针，学生亲手移动箭头来体验“先连新结点，再断旧链接”的顺序要求。",
        ],
        "driving_question": "为什么删除单链表中的一个结点时，通常要先找到它的前驱结点？如果先改错指针，会发生什么？",
        "objectives": [
            "能描述单链表结点由数据域和指针域构成。",
            "能画出带头结点单链表的逻辑结构图。",
            "能完成指定位置插入和删除结点的指针修改，并解释边界情况。",
        ],
        "focus": "结点结构、头结点、前驱结点、插入删除指针顺序。",
        "difficulty_point": "插入时 node.next = p.next 必须先于 p.next = node；删除时必须保存或定位前驱，避免断链。",
        "activities": [
            ("结构搭建", "用卡片和箭头构建 head -> A -> B -> C。", "标注数据域和 next 域。", "能解释结点组成。"),
            ("插入演示", "在 A 与 B 之间插入 X。", "按顺序移动两根箭头。", "不产生断链。"),
            ("删除演示", "删除值为 B 的结点。", "找到前驱 A 并修改 A.next。", "能说明为何需要前驱。"),
            ("代码迁移", "展示 insertAfter 与 deleteByValue。", "补全边界判断。", "头删、空链表能处理。"),
        ],
        "operation_steps": [
            "查找待插入位置的前驱结点 p。",
            "创建新结点 node，并让 node.next 指向 p.next。",
            "再让 p.next 指向 node，完成插入。",
            "删除时找到待删结点 cur 及其前驱 prev。",
            "令 prev.next = cur.next，使链表跨过 cur。",
        ],
        "demo_rows": [
            ["插入前", "head -> A -> B -> C", "p 指向 A", "准备插入 X"],
            ["步骤 1", "X.next = A.next", "X -> B", "新结点先接上后继"],
            ["步骤 2", "A.next = X", "A -> X -> B", "前驱再接新结点"],
            ["删除 B", "X.next = B.next", "X -> C", "跨过待删结点"],
        ],
        "misconceptions": [
            "插入时先写 p.next = node，导致原后继结点丢失。",
            "把“当前结点”当成“前驱结点”，删除时无法修改上一条链接。",
            "忽略空链表、删除首元结点、查找失败等边界情况。",
        ],
        "exercise": "在带头结点单链表中删除值为 x 的第一个结点，写出算法步骤并说明边界情况。",
        "extension": "比较单链表、双链表和数组在插入删除、随机访问上的代价。",
        "assessment": "学生能正确画出插入和删除前后的链表结构，并避免断链或内存泄漏类错误。",
        "code": """insertAfter(p, x):
    node = new Node(x)
    node.next = p.next
    p.next = node

deleteAfter(prev):
    target = prev.next
    if target is not null:
        prev.next = target.next
        release target""",
        "diagram": "linked_list",
    },
    {
        "filename": "ds_bubble_sort_case.docx",
        "title": "数据结构教学案例：冒泡排序过程可视化",
        "topic": "冒泡排序",
        "course": "数据结构",
        "difficulty": "easy",
        "duration": "1 课时 / 45 分钟",
        "keywords": ["冒泡排序", "Bubble Sort", "排序算法", "相邻元素比较", "交换", "时间复杂度"],
        "summary": "通过学生身高排队情境，讲解冒泡排序的相邻比较、交换、轮次收缩、稳定性和 O(n^2) 时间复杂度。",
        "scenario": [
            "体育课按身高从低到高排队。每轮从队首开始比较相邻两名学生，如果前者更高就交换位置，一轮结束后最高的学生会被交换到队尾。",
            "这个“较大元素像气泡一样逐步浮到末尾”的过程，对应冒泡排序 Bubble Sort。课堂中用数字卡片让学生站成一排，每次比较都喊出比较对象和是否交换。",
        ],
        "driving_question": "为什么每完成一轮，最后一个位置就可以不再参与比较？如果一轮中没有发生交换，说明什么？",
        "objectives": [
            "能说出冒泡排序的核心思想：相邻元素比较并交换。",
            "能手工完成一组数据的冒泡排序过程。",
            "能分析冒泡排序的时间复杂度 O(n^2)、稳定性和 swapped 优化。",
        ],
        "focus": "相邻比较、交换、每轮确定最大元素、复杂度分析。",
        "difficulty_point": "理解内层循环边界 n - 1 - i，以及 swapped 标记为何能提前结束。",
        "activities": [
            ("情境导入", "请 5 名学生按 5,3,8,4,2 的数字卡站队。", "观察相邻比较过程。", "能描述比较规则。"),
            ("第一轮演示", "逐对比较并交换。", "记录每次交换后的序列。", "8 到达末尾。"),
            ("轮次收缩", "强调末尾已排好，不再比较。", "标记已确定区域。", "能写出循环边界。"),
            ("优化讨论", "展示已基本有序序列。", "解释 swapped=false。", "知道最好情况 O(n)。"),
        ],
        "operation_steps": [
            "外层循环控制排序轮次。",
            "内层循环从左到右比较相邻元素。",
            "若 a[j] > a[j+1]，交换二者。",
            "每轮结束后，当前未排序区最大元素固定到末尾。",
            "如果某轮没有交换，可提前结束。",
        ],
        "demo_rows": [
            ["初始", "5 3 8 4 2", "未排序", "准备第一轮"],
            ["比较 5 和 3", "3 5 8 4 2", "交换", "较大值右移"],
            ["比较 8 和 4", "3 5 4 8 2", "交换", "8 继续右移"],
            ["比较 8 和 2", "3 5 4 2 8", "交换", "8 固定到末尾"],
            ["第二轮后", "3 4 2 5 8", "5 固定", "未排序区缩小"],
        ],
        "misconceptions": [
            "把每轮最大元素固定到末尾误写成最小元素固定到开头。",
            "内层循环没有减去 i，导致已排序区域被重复比较。",
            "认为冒泡排序总是适合大规模数据，忽略 O(n^2) 代价。",
        ],
        "exercise": "对数组 [6, 1, 4, 2, 5] 执行冒泡排序，写出每一轮排序后的数组。",
        "extension": "比较冒泡排序与选择排序、插入排序在交换次数和适用场景上的差异。",
        "assessment": "学生能准确展示比较、交换、轮次终止条件，并能说明冒泡排序适合用于教学入门而非大规模数据排序。",
        "code": """for i from 0 to n - 2:
    swapped = false
    for j from 0 to n - 2 - i:
        if a[j] > a[j + 1]:
            swap(a[j], a[j + 1])
            swapped = true
    if not swapped: break""",
        "diagram": "bubble_sort",
    },
    {
        "filename": "ds_binary_search_case.docx",
        "title": "数据结构教学案例：折半查找的有序性条件",
        "topic": "折半查找",
        "course": "数据结构",
        "difficulty": "medium",
        "duration": "1 课时 / 45 分钟",
        "keywords": ["折半查找", "二分查找", "Binary Search", "有序表", "low", "high", "mid"],
        "summary": "以图书馆按编号查书为情境，讲解折半查找的有序性前提、low/high/mid 区间收缩和 O(log n) 复杂度。",
        "scenario": [
            "图书馆书架按编号递增排列。查找某本书时，不需要从第一本开始顺序扫描，而是先查看中间编号，再判断目标在左半区还是右半区。",
            "教师给出 1 到 31 的编号卡片，让学生扮演查找者，每次只允许翻开中间卡片，并用 low、high、mid 三个磁贴标记当前查找区间。",
        ],
        "driving_question": "折半查找为什么必须要求数据有序？如果数组没有排序，mid 的比较还能排除一半数据吗？",
        "objectives": [
            "能说明折半查找必须作用于有序顺序表。",
            "能正确更新 low、high、mid 三个变量。",
            "能分析折半查找时间复杂度 O(log n)，并与顺序查找比较。",
        ],
        "focus": "有序性前提、区间收缩、low/high/mid 更新。",
        "difficulty_point": "target 与 a[mid] 比较后的区间更新方向，以及 mid = low + (high - low) / 2 的安全写法。",
        "activities": [
            ("情境建模", "展示按编号排列的书架。", "指出有序性带来的排除能力。", "能说出“必须有序”。"),
            ("手工查找", "用 1 到 31 查找 23。", "记录 low、high、mid。", "每轮区间缩小一半。"),
            ("边界讨论", "比较 low <= high 与 low < high。", "分析找不到元素的退出条件。", "不会漏查最后一个元素。"),
            ("复杂度对比", "画出查找次数随 n 增长的表格。", "比较顺序查找和折半查找。", "能解释 O(log n)。"),
        ],
        "operation_steps": [
            "令 low 指向首元素，high 指向末元素。",
            "当 low <= high 时，计算 mid = low + (high - low) / 2。",
            "若 a[mid] 等于目标，查找成功。",
            "若 a[mid] 小于目标，说明目标只可能在右半区，更新 low = mid + 1。",
            "若 a[mid] 大于目标，说明目标只可能在左半区，更新 high = mid - 1。",
        ],
        "demo_rows": [
            ["1", "1", "31", "16", "23 > 16，去右半区"],
            ["2", "17", "31", "24", "23 < 24，去左半区"],
            ["3", "17", "23", "20", "23 > 20，去右半区"],
            ["4", "21", "23", "22", "23 > 22，去右半区"],
            ["5", "23", "23", "23", "查找成功"],
        ],
        "misconceptions": [
            "在无序数组上使用折半查找，导致错误排除数据。",
            "mid 更新后没有 +1 或 -1，区间无法缩小，可能死循环。",
            "使用 (low + high) / 2 时忽略整数溢出风险。",
        ],
        "exercise": "在有序数组 [3, 7, 12, 18, 25, 31, 40] 中查找 25，写出 low、high、mid 的变化。",
        "extension": "讨论二分思想在答案空间搜索、查找第一个满足条件的位置中的应用。",
        "assessment": "学生能明确有序性前提，并能正确控制查找区间，即视为达标。",
        "code": """while low <= high:
    mid = low + (high - low) / 2
    if a[mid] == key: return mid
    if a[mid] < key:
        low = mid + 1
    else:
        high = mid - 1

return -1""",
        "diagram": "binary_search",
    },
    {
        "filename": "ds_binary_tree_traversal_case.docx",
        "title": "数据结构教学案例：二叉树的三种遍历",
        "topic": "二叉树",
        "course": "数据结构",
        "difficulty": "medium",
        "duration": "2 课时 / 90 分钟",
        "keywords": ["二叉树", "Binary Tree", "先序遍历", "中序遍历", "后序遍历", "递归", "树结构"],
        "summary": "以组织结构图为情境，讲解二叉树的结点、左右子树，以及先序、中序、后序遍历规则。",
        "scenario": [
            "公司组织结构可以抽象成树结构。若每个部门最多拆分为两个下级部门，就可以用二叉树表示。访问整棵二叉树时，根结点、左子树、右子树的访问顺序不同，会得到不同遍历结果。",
            "课堂使用 A(B(D,E), C(F,G)) 这棵树，让学生用彩色笔标出根、左子树、右子树，再按“根左右”“左根右”“左右根”三套规则走完整棵树。",
        ],
        "driving_question": "同一棵二叉树为什么会有不同遍历序列？递归访问时，什么时候访问根结点？",
        "objectives": [
            "能描述二叉树中根结点、左子树、右子树的关系。",
            "能区分先序遍历、中序遍历和后序遍历。",
            "能根据给定二叉树写出三种遍历序列，并解释递归过程。",
        ],
        "focus": "根、左子树、右子树的访问顺序；递归分解。",
        "difficulty_point": "每进入一棵子树都要重复同样规则，而不是只在第一层应用遍历顺序。",
        "activities": [
            ("结构识别", "给出树 A(B(D,E), C(F,G))。", "标出每个结点的左右孩子。", "能说出根和子树。"),
            ("先序遍历", "强调根、左、右。", "按规则读出 A B D E C F G。", "根先访问。"),
            ("中序遍历", "强调左、根、右。", "读出 D B E A F C G。", "根在左右之间。"),
            ("后序遍历", "强调左、右、根。", "读出 D E B F G C A。", "根最后访问。"),
            ("递归迁移", "展示遍历伪代码。", "在子树上重复规则。", "能解释递归终止。"),
        ],
        "operation_steps": [
            "先序遍历：先访问根结点，再遍历左子树，最后遍历右子树。",
            "中序遍历：先遍历左子树，再访问根结点，最后遍历右子树。",
            "后序遍历：先遍历左子树，再遍历右子树，最后访问根结点。",
            "遇到空子树时递归返回。",
        ],
        "demo_rows": [
            ["先序", "根 -> 左 -> 右", "A B D E C F G", "根结点最早出现"],
            ["中序", "左 -> 根 -> 右", "D B E A F C G", "常用于二叉搜索树排序输出"],
            ["后序", "左 -> 右 -> 根", "D E B F G C A", "适合先处理子结点再处理父结点"],
        ],
        "misconceptions": [
            "只在第一层套用遍历规则，进入子树后改成从左到右扫描。",
            "把中序遍历误认为一定得到升序，忽略只有二叉搜索树才成立。",
            "递归终止条件缺失，空结点仍继续访问。",
        ],
        "exercise": "对二叉树 A(B(D,E), C(,F)) 分别写出先序、中序、后序遍历序列。",
        "extension": "引出二叉搜索树、表达式树和非递归遍历中栈的使用。",
        "assessment": "学生能在遍历过程中保持递归访问规则一致，并能解释三种遍历的差异。",
        "code": """preOrder(root):
    if root is null: return
    visit(root)
    preOrder(root.left)
    preOrder(root.right)

inOrder(root):
    if root is null: return
    inOrder(root.left)
    visit(root)
    inOrder(root.right)""",
        "diagram": "binary_tree",
    },
    {
        "filename": "ds_graph_bfs_case.docx",
        "title": "数据结构教学案例：图的广度优先搜索",
        "topic": "图的广度优先搜索",
        "course": "数据结构",
        "difficulty": "hard",
        "duration": "2 课时 / 90 分钟",
        "keywords": ["图", "Graph", "广度优先搜索", "BFS", "队列", "邻接表", "最短路径"],
        "summary": "以校园地图最少步数问题为情境，讲解图的邻接表表示、BFS 的队列过程和无权图最短路径层次含义。",
        "scenario": [
            "校园内多个教学楼通过道路连接。若每条道路长度近似相同，从图书馆出发寻找最少经过几条道路到达实验楼，可以使用广度优先搜索 BFS。",
            "教师把教学楼抽象为顶点，道路抽象为边。学生从起点 A 出发，按“先访问距离 1 的全部邻居，再访问距离 2 的顶点”的层次规则推进。",
        ],
        "driving_question": "BFS 为什么需要队列？它为什么能在无权图中找到从起点到其他顶点的最短边数？",
        "objectives": [
            "能将校园道路抽象为图的顶点和边。",
            "能使用邻接表表示无向图。",
            "能说明 BFS 需要队列保存待访问顶点，并按层次扩展。",
            "能根据 BFS 过程记录 visited、parent 和 distance。",
        ],
        "focus": "图抽象、邻接表、队列、按层访问、无权图最短路径。",
        "difficulty_point": "区分“入队时标记 visited”和“出队时访问邻居”，避免同一顶点被重复入队。",
        "activities": [
            ("图建模", "将教学楼 A、B、C、D、E、F 画成无向图。", "根据道路写邻接表。", "能把现实关系转成边。"),
            ("BFS 演示", "从 A 出发维护队列。", "记录队列变化和 visited。", "访问顺序正确。"),
            ("层次理解", "标注每个顶点到 A 的距离。", "解释为什么先到达就是最短边数。", "能说出无权图前提。"),
            ("路径还原", "记录 parent 数组。", "还原 A 到 F 的路径。", "能从终点倒推父结点。"),
        ],
        "operation_steps": [
            "将起点入队，并立即标记 visited。",
            "当队列不为空时，取出队头顶点 v。",
            "遍历 v 的所有邻接点 u。",
            "若 u 未访问，则标记 visited，记录 parent[u] = v，并将 u 入队。",
            "队列保证先发现的低层顶点先扩展，因此形成按层遍历。",
        ],
        "demo_rows": [
            ["初始", "A", "A", "dist[A]=0"],
            ["访问 A", "B, C", "A, B, C", "B/C 距离为 1"],
            ["访问 B", "C, D", "A, B, C, D", "D 距离为 2"],
            ["访问 C", "D, E", "A, B, C, D, E", "E 距离为 2"],
            ["访问 D/E", "F", "A, B, C, D, E, F", "F 距离为 3"],
        ],
        "misconceptions": [
            "把 BFS 写成递归深挖，实际变成 DFS。",
            "出队后才标记 visited，导致同一顶点被多个父结点重复入队。",
            "认为 BFS 能处理带权最短路径，忽略无权或等权前提。",
        ],
        "exercise": "给定邻接表 A: B,C; B: A,D; C: A,E; D: B,F; E: C,F; F: D,E，从 A 开始写出 BFS 访问序列和 parent 数组。",
        "extension": "对比 BFS 与 DFS，并引出 Dijkstra 算法处理带权最短路径的必要性。",
        "assessment": "学生能正确使用队列推进 BFS，并能解释层次遍历与最短路径的关系。",
        "code": """BFS(start):
    queue = new Queue()
    visited[start] = true
    distance[start] = 0
    enqueue(start)

    while queue is not empty:
        v = dequeue()
        for each neighbor u of v:
            if not visited[u]:
                visited[u] = true
                parent[u] = v
                distance[u] = distance[v] + 1
                enqueue(u)""",
        "diagram": "graph_bfs",
    },
]


def load_font(size, bold=False):
    for path in FONT_CANDIDATES:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                pass
    return ImageFont.load_default()


FONT_TITLE = load_font(36)
FONT_LABEL = load_font(26)
FONT_SMALL = load_font(22)
FONT_CODE = load_font(20)


def draw_arrow(draw, start, end, color="#475467", width=4):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    spread = 0.55
    p1 = (
        end[0] - length * math.cos(angle - spread),
        end[1] - length * math.sin(angle - spread),
    )
    p2 = (
        end[0] - length * math.cos(angle + spread),
        end[1] - length * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill=color)


def centered_text(draw, box, text, font, fill="#101828"):
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2 - 2), text, font=font, fill=fill)


def label(draw, xy, text, fill="#344054", font=FONT_SMALL):
    draw.text(xy, text, font=font, fill=fill)


def new_canvas(title):
    img = Image.new("RGB", (1400, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((40, 36, 1360, 724), radius=28, fill="#FFFFFF", outline="#EAECF0", width=3)
    draw.text((70, 58), title, font=FONT_TITLE, fill="#0B2545")
    return img, draw


def draw_queue(path):
    img, draw = new_canvas("队列 Queue：从 rear 入队，从 front 出队")
    x0, y0, cell_w, cell_h = 190, 270, 170, 120
    items = ["A", "B", "C", "", ""]
    for i, item in enumerate(items):
        x = x0 + i * cell_w
        fill = "#EAF2FF" if item else "#F8FAFC"
        draw.rounded_rectangle((x, y0, x + cell_w, y0 + cell_h), radius=14, fill=fill, outline="#98A2B3", width=3)
        centered_text(draw, (x, y0, x + cell_w, y0 + cell_h), item or "空", FONT_LABEL, "#0B2545" if item else "#98A2B3")
        centered_text(draw, (x, y0 + cell_h + 16, x + cell_w, y0 + cell_h + 56), str(i), FONT_SMALL, "#667085")
    draw_arrow(draw, (130, y0 + 60), (x0 - 20, y0 + 60), "#16A34A", 5)
    label(draw, (72, y0 + 90), "dequeue\n队头 front", "#166534")
    draw_arrow(draw, (x0 + 5 * cell_w + 90, y0 + 60), (x0 + 5 * cell_w + 5, y0 + 60), "#EA580C", 5)
    label(draw, (x0 + 5 * cell_w + 104, y0 + 34), "enqueue\n队尾 rear", "#9A3412")
    label(draw, (190, 520), "证据片段：队列遵循 FIFO，先进入的 A 会先从 front 出队。", "#344054")
    img.save(path)


def draw_circular_queue(path):
    img, draw = new_canvas("循环队列 Circular Queue：取模让指针回绕")
    cx, cy, r = 700, 400, 210
    positions = []
    for i in range(6):
        angle = -math.pi / 2 + i * 2 * math.pi / 6
        x = cx + r * math.cos(angle)
        y = cy + r * math.sin(angle)
        positions.append((x, y))
        fill = "#EAF2FF" if i in [2, 3, 4, 0] else "#F8FAFC"
        draw.ellipse((x - 58, y - 58, x + 58, y + 58), fill=fill, outline="#98A2B3", width=3)
        centered_text(draw, (x - 58, y - 40, x + 58, y + 22), str(i), FONT_LABEL)
    draw.ellipse((cx - r - 38, cy - r - 38, cx + r + 38, cy + r + 38), outline="#D0D5DD", width=4)
    label(draw, (210, 210), "front = 2\n指向队头元素", "#166534")
    draw_arrow(draw, (340, 260), positions[2], "#16A34A", 5)
    label(draw, (930, 210), "rear = 1\n指向下一个空位", "#9A3412")
    draw_arrow(draw, (930, 278), positions[1], "#EA580C", 5)
    label(draw, (350, 615), "核心公式：rear = (rear + 1) % maxSize", "#344054")
    img.save(path)


def draw_stack(path):
    img, draw = new_canvas("栈 Stack：最近入栈的左括号最先匹配")
    x, y, w, h = 560, 210, 300, 88
    items = ["{", "[", "("]
    for i, item in enumerate(items):
        yy = y + (len(items) - 1 - i) * h
        draw.rounded_rectangle((x, yy, x + w, yy + h), radius=12, fill="#EAF2FF", outline="#98A2B3", width=3)
        centered_text(draw, (x, yy, x + w, yy + h), item, FONT_TITLE, "#0B2545")
    draw.rounded_rectangle((x, y + len(items) * h, x + w, y + len(items) * h + 24), radius=8, fill="#D0D5DD")
    draw_arrow(draw, (1040, y + 44), (x + w + 18, y + 44), "#EA580C", 5)
    label(draw, (1050, y + 70), "pop\n遇到右括号", "#9A3412")
    draw_arrow(draw, (380, y + 44), (x - 18, y + 44), "#16A34A", 5)
    label(draw, (250, y + 70), "push\n遇到左括号", "#166534")
    label(draw, (390, 610), "LIFO：后入栈的 '(' 会先与 ')' 匹配。", "#344054")
    img.save(path)


def draw_linked_list(path):
    img, draw = new_canvas("单链表 Linked List：插入只修改相邻指针")
    nodes = [("head", 120), ("A", 340), ("X", 560), ("B", 780), ("C", 1000)]
    y = 330
    for name, x in nodes:
        draw.rounded_rectangle((x, y, x + 140, y + 90), radius=14, fill="#EAF2FF" if name != "X" else "#ECFDF3", outline="#98A2B3", width=3)
        draw.line((x + 92, y, x + 92, y + 90), fill="#98A2B3", width=3)
        centered_text(draw, (x, y, x + 92, y + 90), name, FONT_LABEL)
        centered_text(draw, (x + 92, y, x + 140, y + 90), "next", FONT_SMALL, "#667085")
    for (_, x1), (_, x2) in zip(nodes, nodes[1:]):
        draw_arrow(draw, (x1 + 140, y + 45), (x2 - 12, y + 45), "#475467", 4)
    label(draw, (360, 235), "插入 X 的顺序：1. X.next = A.next    2. A.next = X", "#344054")
    label(draw, (456, 500), "如果先让 A.next = X，原来的 B 可能丢失，形成断链。", "#9A3412")
    img.save(path)


def draw_bubble_sort(path):
    img, draw = new_canvas("冒泡排序 Bubble Sort：相邻比较，较大元素右移")
    rows = [
        ("初始", [5, 3, 8, 4, 2], None),
        ("比较 5/3", [3, 5, 8, 4, 2], (0, 1)),
        ("比较 8/4", [3, 5, 4, 8, 2], (2, 3)),
        ("比较 8/2", [3, 5, 4, 2, 8], (3, 4)),
    ]
    x0, y0, cell_w, cell_h = 300, 170, 110, 66
    for r, (label_text, nums, pair) in enumerate(rows):
        y = y0 + r * 115
        label(draw, (110, y + 16), label_text, "#344054")
        for i, n in enumerate(nums):
            x = x0 + i * cell_w
            fill = "#FFF4E5" if pair and i in pair else "#EAF2FF"
            if r == len(rows) - 1 and i == 4:
                fill = "#ECFDF3"
            draw.rounded_rectangle((x, y, x + 88, y + cell_h), radius=10, fill=fill, outline="#98A2B3", width=3)
            centered_text(draw, (x, y, x + 88, y + cell_h), str(n), FONT_LABEL)
        if pair:
            px = x0 + pair[0] * cell_w + 44
            qx = x0 + pair[1] * cell_w + 44
            draw.arc((px, y - 34, qx, y + 40), 180, 360, fill="#EA580C", width=4)
    label(draw, (820, 515), "第一轮结束：8 已固定在末尾，下一轮比较范围缩小。", "#344054")
    img.save(path)


def draw_binary_search(path):
    img, draw = new_canvas("折半查找 Binary Search：每次排除一半区间")
    nums = [1, 3, 7, 12, 16, 20, 22, 23, 24, 28, 31]
    x0, y0, cell_w, cell_h = 130, 270, 110, 82
    for i, n in enumerate(nums):
        x = x0 + i * cell_w
        fill = "#EAF2FF"
        if n == 23:
            fill = "#ECFDF3"
        if n in [16, 24, 20, 22]:
            fill = "#FFF4E5"
        draw.rounded_rectangle((x, y0, x + 90, y0 + cell_h), radius=10, fill=fill, outline="#98A2B3", width=3)
        centered_text(draw, (x, y0, x + 90, y0 + cell_h), str(n), FONT_LABEL)
    draw_arrow(draw, (x0 + 4 * cell_w + 45, 215), (x0 + 4 * cell_w + 45, y0 - 8), "#EA580C", 5)
    label(draw, (x0 + 4 * cell_w - 44, 178), "mid=16", "#9A3412")
    draw_arrow(draw, (x0 + 7 * cell_w + 45, 470), (x0 + 7 * cell_w + 45, y0 + cell_h + 8), "#16A34A", 5)
    label(draw, (x0 + 7 * cell_w - 50, 500), "目标 23", "#166534")
    label(draw, (260, 590), "23 > 16，左半区可整体排除；有序性是排除一半数据的前提。", "#344054")
    img.save(path)


def draw_binary_tree(path):
    img, draw = new_canvas("二叉树 Binary Tree：三种遍历关注根的位置")
    pos = {
        "A": (700, 165),
        "B": (430, 320),
        "C": (970, 320),
        "D": (300, 500),
        "E": (560, 500),
        "F": (840, 500),
        "G": (1100, 500),
    }
    edges = [("A", "B"), ("A", "C"), ("B", "D"), ("B", "E"), ("C", "F"), ("C", "G")]
    for a, b in edges:
        draw.line([pos[a], pos[b]], fill="#98A2B3", width=5)
    for name, (x, y) in pos.items():
        draw.ellipse((x - 54, y - 54, x + 54, y + 54), fill="#EAF2FF", outline="#2563EB", width=4)
        centered_text(draw, (x - 54, y - 54, x + 54, y + 54), name, FONT_TITLE, "#0B2545")
    label(draw, (120, 622), "先序：A B D E C F G    中序：D B E A F C G    后序：D E B F G C A", "#344054")
    img.save(path)


def draw_graph_bfs(path):
    img, draw = new_canvas("图的 BFS：队列保证按层扩展")
    pos = {
        "A": (260, 380),
        "B": (490, 230),
        "C": (490, 530),
        "D": (760, 230),
        "E": (760, 530),
        "F": (1030, 380),
    }
    edges = [("A", "B"), ("A", "C"), ("B", "D"), ("C", "E"), ("D", "F"), ("E", "F")]
    for a, b in edges:
        draw.line([pos[a], pos[b]], fill="#98A2B3", width=5)
    levels = {"A": "#ECFDF3", "B": "#EAF2FF", "C": "#EAF2FF", "D": "#FFF4E5", "E": "#FFF4E5", "F": "#FEF3F2"}
    for name, (x, y) in pos.items():
        draw.ellipse((x - 58, y - 58, x + 58, y + 58), fill=levels[name], outline="#475467", width=4)
        centered_text(draw, (x - 58, y - 58, x + 58, y + 58), name, FONT_TITLE, "#0B2545")
    label(draw, (145, 608), "层次：A(0) -> B,C(1) -> D,E(2) -> F(3)；队列保存下一批待扩展顶点。", "#344054")
    img.save(path)


DIAGRAM_DRAWERS = {
    "queue": draw_queue,
    "circular_queue": draw_circular_queue,
    "stack": draw_stack,
    "linked_list": draw_linked_list,
    "bubble_sort": draw_bubble_sort,
    "binary_search": draw_binary_search,
    "binary_tree": draw_binary_tree,
    "graph_bfs": draw_graph_bfs,
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 15, ACCENT_DARK, 13, 6),
        ("Heading 2", 12.5, ACCENT, 9, 4),
        ("Heading 3", 11.5, ACCENT_DARK, 6, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def add_title(doc, case):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(case["title"])
    set_run_font(run, 20, True, ACCENT_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("测试数据用途：用于教学案例库精准匹配、证据片段检索与教案生成案例匹配评估")
    set_run_font(r, 10, False, MUTED)


def style_paragraph_runs(paragraph, size=10.5, color=None, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_field_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    first_row = True
    for label_text, value in rows:
        row = table.add_row()
        if first_row:
            set_repeat_table_header(row)
            first_row = False
        cells = row.cells
        cells[0].text = label_text
        cells[1].text = value
        set_cell_width(cells[0], 2100)
        set_cell_width(cells[1], 7600)
        set_cell_shading(cells[0], "EEF4FF")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                style_paragraph_runs(p)
        cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body, fill=SOFT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "C7D7FE")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9700)
    set_cell_margins(cell, top=170, bottom=170, start=180, end=180)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(title + "：")
    set_run_font(r1, 10.5, True, ACCENT_DARK)
    r2 = p.add_run(body)
    set_run_font(r2, 10.5, False, ACCENT_DARK)
    doc.add_paragraph()


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "D0D5DD")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT_GRAY)
    set_cell_width(cell, 9700)
    set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    set_run_font(run, 9.2, False, "1D2939", "Consolas")
    doc.add_paragraph()


def add_table(doc, headers, rows, widths, header_fill="EEF4FF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph_runs(p, size=9.6, bold=True, color=ACCENT_DARK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = str(value)
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                if widths[idx] <= 1350:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style_paragraph_runs(p, size=9.2)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.45))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", caption)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    r = caption_p.add_run(caption)
    set_run_font(r, 9, False, MUTED)


def generate_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for case in CASES:
        image_path = ASSET_DIR / case["filename"].replace(".docx", ".png")
        DIAGRAM_DRAWERS[case["diagram"]](image_path)


def build_doc(case):
    doc = Document()
    configure_styles(doc)
    add_title(doc, case)

    add_field_table(
        doc,
        [
            ("适用课程", case["course"]),
            ("核心课题", case["topic"]),
            ("难度等级", case["difficulty"]),
            ("建议课时", case["duration"]),
            ("知识点标签", "、".join(case["keywords"])),
            ("案例摘要", case["summary"]),
        ],
    )

    doc.add_heading("一、教学情境与问题驱动", level=1)
    for para in case["scenario"]:
        doc.add_paragraph(para)
    add_callout(doc, "驱动问题", case["driving_question"], SOFT_ORANGE)

    image_path = ASSET_DIR / case["filename"].replace(".docx", ".png")
    add_figure(doc, image_path, f"图 1  {case['topic']}核心结构或过程示意图")

    doc.add_heading("二、教学目标", level=1)
    add_bullets(doc, case["objectives"])

    doc.add_heading("三、教学重点与难点", level=1)
    add_field_table(
        doc,
        [
            ("教学重点", case["focus"]),
            ("教学难点", case["difficulty_point"]),
            ("检索锚点", f"本案例应围绕“{case['topic']}”命中，证据优先来自知识点标签、核心算法、课堂活动和示意图说明。"),
        ],
    )

    doc.add_heading("四、课堂活动设计", level=1)
    add_table(
        doc,
        ["环节", "教师活动", "学生活动", "观察证据"],
        case["activities"],
        [1150, 3300, 3000, 2250],
    )

    doc.add_heading("五、核心算法或操作过程", level=1)
    doc.add_paragraph(f"以下片段用于突出本案例的核心知识点“{case['topic']}”，也是案例检索时应命中的主要证据。")
    add_numbered(doc, case["operation_steps"])
    add_code_block(doc, case["code"])

    doc.add_heading("六、过程演示与板书记录", level=1)
    first_header = {
        "queue": ["操作", "队列状态", "front", "rear", "说明"],
        "circular_queue": ["操作", "指针状态", "数组状态", "说明"],
        "stack": ["扫描字符", "动作", "栈状态", "说明"],
        "linked_list": ["阶段", "指针语句", "链表状态", "说明"],
        "bubble_sort": ["阶段", "序列", "动作", "说明"],
        "binary_search": ["轮次", "low", "high", "mid", "判断"],
        "binary_tree": ["遍历方式", "访问规则", "遍历序列", "教学提示"],
        "graph_bfs": ["阶段", "队列", "已访问", "说明"],
    }[case["diagram"]]
    widths_by_cols = {
        3: [1800, 2800, 3900],
        4: [1450, 2600, 3200, 2450],
        5: [1100, 1600, 1600, 1600, 3900],
    }
    add_table(doc, first_header, case["demo_rows"], widths_by_cols[len(first_header)], "ECFDF3")

    doc.add_heading("七、易错点诊断与教师追问", level=1)
    add_bullets(doc, case["misconceptions"])
    add_callout(doc, "追问方式", "让学生先画状态图，再解释变量或指针为什么这样变化；若只会背代码，要求回到示意图重新说明。", SOFT_RED)

    doc.add_heading("八、课堂练习与拓展", level=1)
    add_field_table(
        doc,
        [
            ("基础练习", case["exercise"]),
            ("拓展任务", case["extension"]),
        ],
    )

    doc.add_heading("九、评价方式", level=1)
    doc.add_paragraph(case["assessment"])
    add_table(
        doc,
        ["评价维度", "达标表现", "需要支持的表现"],
        [
            ["概念理解", f"能用自己的话解释“{case['topic']}”及其适用条件。", "只能复述名称，不能联系情境或示意图。"],
            ["过程表达", "能画出关键状态变化，并说明每一步变量、指针或队列变化。", "状态跳步，边界条件遗漏。"],
            ["代码迁移", "能把课堂流程转写为伪代码或程序片段。", "语句顺序错误，缺少异常判断。"],
        ],
        [1900, 4800, 3000],
        "FFF4E5",
    )

    doc.add_heading("十、检索测试提示", level=1)
    doc.add_paragraph(
        f"本案例应在用户输入“数据结构 / {case['topic']}”时被优先匹配。"
        "如果用户输入其他不相关课题，除非本文明确包含该课题证据，否则不应被标记为精准匹配。"
    )
    doc.add_paragraph("建议系统返回证据片段时优先引用“知识点标签”“核心算法或操作过程”“课堂活动设计”“过程演示与板书记录”中的句子。")
    doc.add_paragraph("负向测试建议：输入相邻但不同的课题时，系统应能区分本案例的核心证据与泛化词。例如“栈”不应误匹配为普通“队列”案例。")

    path = OUT_DIR / case["filename"]
    doc.save(path)
    return path


def write_manifest(generated):
    manifest_path = OUT_DIR / "manifest.csv"
    with manifest_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["file", "title", "category", "difficulty", "courseName", "topic", "keywords"])
        for case, path in generated:
            writer.writerow(
                [
                    path.name,
                    case["title"],
                    "course_design",
                    case["difficulty"],
                    case["course"],
                    case["topic"],
                    "、".join(case["keywords"]),
                ]
            )
    return manifest_path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    RENDERED_DIR.mkdir(parents=True, exist_ok=True)
    generate_diagrams()
    generated = []
    for case in CASES:
        generated.append((case, build_doc(case)))
    manifest_path = write_manifest(generated)

    print(str(OUT_DIR))
    for _, path in generated:
        print(path.name)
    print(manifest_path.name)
    print(str(ASSET_DIR))


if __name__ == "__main__":
    main()
